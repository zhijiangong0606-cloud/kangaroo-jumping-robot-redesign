# -*- coding: utf-8 -*-
"""把实验报告 Markdown 转成 Word(.docx)，处理标题/表格/段落，并在合适位置嵌图。"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"C:\Users\Gzj\Desktop\机械设计实验"
MD = os.path.join(BASE, "袋鼠仿生弹跳机器人实验报告.md")
OUT = os.path.join(BASE, "袋鼠仿生弹跳机器人实验报告.docx")
IMGDIR = os.path.join(BASE, "渲染图")

doc = Document()

# ---- 基础样式：中文宋体 / 标题黑体 ----
def set_cjk(run, font="宋体"):
    run.font.name = font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 行距
from docx.oxml.ns import qn as _q
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)

# 图片：标题词 -> (文件, 图注) ，插在对应小节末尾
SECTION_IMAGES = {
    "五、双模式仿生尾的设计": [
        ("tail_dualmode_iso.png", "图1  双模式尾等轴测对比：左为蓄能位（尾下垂落地支撑），右为起跳位（尾上扬配重）"),
        ("tail_dualmode_front.png", "图2  双模式尾前视对比：蓄能位 / 起跳位"),
    ],
    "七、验证": [
        ("装配工程图_预览.png", "图3  装配工程图（前视/俯视/右视 + 等轴测，含最新双模式尾）"),
    ],
}

def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    sizes = {1: 18, 2: 15, 3: 13}
    run.font.size = Pt(sizes.get(level, 12))
    set_cjk(run, "黑体")
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_para(text):
    p = doc.add_paragraph()
    # 处理行内 **粗体**
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); run.bold = True
        else:
            run = p.add_run(part)
        set_cjk(run, "宋体")
    return p

def add_image(fn, caption):
    path = os.path.join(IMGDIR, fn)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size = Pt(10.5); r.italic = True
    set_cjk(r, "宋体")

print("converter ready")

def add_table(rows):
    # rows: list of list[str]，第一行为表头
    ncol = len(rows[0])
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ""
            val = val.replace('**','')
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(val)
            set_cjk(run, "宋体")
            run.font.size = Pt(10.5)
            if ri == 0:
                run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

with open(MD, encoding='utf-8') as f:
    lines = f.readlines()

current_section = None
i = 0
pending_table = []

def flush_table():
    global pending_table
    if pending_table:
        # 过滤分隔行 |---|---|
        clean = [r for r in pending_table if not re.match(r'^\s*\|?[\s:\-\|]+\|?\s*$', '|'.join(r))]
        if clean:
            add_table(clean)
    pending_table = []

def emit_section_images(sec):
    if sec in SECTION_IMAGES:
        for fn, cap in SECTION_IMAGES[sec]:
            add_image(fn, cap)

while i < len(lines):
    raw = lines[i].rstrip('\n')
    line = raw.strip()

    # 表格行
    if line.startswith('|') and line.endswith('|'):
        cells = [c.strip() for c in line.strip('|').split('|')]
        pending_table.append(cells)
        i += 1
        continue
    else:
        if pending_table:
            flush_table()

    if not line:
        i += 1
        continue

    m = re.match(r'^(#{1,3})\s+(.*)', line)
    if m:
        # 进入新小节前，先把上一节的图补上
        emit_section_images(current_section)
        level = len(m.group(1))
        title = m.group(2).strip()
        add_heading(title, level)
        if level == 2:
            current_section = title
        elif level == 1:
            current_section = None
        i += 1
        continue

    # 列表项
    if re.match(r'^[-*]\s+', line):
        text = re.sub(r'^[-*]\s+', '', line)
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.8)
        parts = re.split(r'(\*\*.+?\*\*)', text)
        bullet = p.add_run("• "); set_cjk(bullet, "宋体")
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2]); run.bold = True
            else:
                run = p.add_run(part)
            set_cjk(run, "宋体")
        i += 1
        continue

    # 普通段落
    add_para(line)
    i += 1

# 文末把最后一节的图补上
flush_table()
emit_section_images(current_section)

doc.save(OUT)
print("SAVED:", OUT)
print("段落数:", len(doc.paragraphs), "表格数:", len(doc.tables))
